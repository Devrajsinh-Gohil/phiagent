import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredMarkdownLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredFileLoader,
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles loading and processing of various document types for RAG."""
    
    @staticmethod
    def load_document(file_path: str) -> str:
        """Load document content based on file extension."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return DocumentProcessor._load_pdf(file_path)
            elif file_ext == '.docx':
                return DocumentProcessor._load_docx(file_path)
            elif file_ext == '.txt':
                return DocumentProcessor._load_text(file_path)
            elif file_ext in ['.md', '.markdown']:
                return DocumentProcessor._load_markdown(file_path)
            else:
                return DocumentProcessor._load_generic(file_path)
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _load_pdf(file_path: Path) -> str:
        """Load text from PDF file."""
        reader = PdfReader(file_path)
        text = "\n".join([page.extract_text() for page in reader.pages])
        return text
    
    @staticmethod
    def _load_docx(file_path: Path) -> str:
        """Load text from Word document."""
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    @staticmethod
    def _load_text(file_path: Path) -> str:
        """Load text from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def _load_markdown(file_path: Path) -> str:
        """Load text from markdown file."""
        loader = UnstructuredMarkdownLoader(str(file_path))
        return loader.load()[0].page_content
    
    @staticmethod
    def _load_generic(file_path: Path) -> str:
        """Fallback method for other file types."""
        loader = UnstructuredFileLoader(str(file_path))
        return loader.load()[0].page_content
    
    @staticmethod
    def split_documents(
        documents: List[LangchainDocument], 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200
    ) -> List[LangchainDocument]:
        """Split documents into chunks for processing."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            add_start_index=True,
        )
        
        return text_splitter.split_documents(documents)
    
    @staticmethod
    def process_directory(
        directory: str, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200
    ) -> List[LangchainDocument]:
        """Process all supported documents in a directory."""
        directory = Path(directory)
        if not directory.is_dir():
            raise NotADirectoryError(f"Directory not found: {directory}")
        
        supported_extensions = ['.pdf', '.docx', '.txt', '.md', '.markdown']
        documents = []
        
        for ext in supported_extensions:
            for file_path in directory.glob(f"*{ext}"):
                try:
                    content = DocumentProcessor.load_document(str(file_path))
                    doc = LangchainDocument(
                        page_content=content,
                        metadata={
                            "source": str(file_path.name),
                            "file_path": str(file_path),
                        }
                    )
                    documents.append(doc)
                    logger.info(f"Processed document: {file_path.name}")
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {str(e)}")
        
        # Split documents into chunks
        if documents:
            return DocumentProcessor.split_documents(documents, chunk_size, chunk_overlap)
        return []

# Example usage:
if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    docs = processor.process_directory("path/to/your/documents")
    print(f"Processed {len(docs)} document chunks")
